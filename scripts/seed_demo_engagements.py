"""Seed two synthetic E2E demo engagements through the live application API.

The script deliberately uses the same HTTP routes, ingestion pipeline, interview
agent, transcript finalization, retrieval, and Deep Agent insight flow as the UI.
Only the passage of time for the expired-session example is simulated directly
in SQLite; the draft conversation itself is produced by the live interview agent.
"""

# The seed keeps realistic interview answers as readable source literals, uses
# python-docx's documented OOXML escape hatch, and raises script-local failures.
# ruff: noqa: E501, PLR2004, SLF001, TRY003

from __future__ import annotations

import argparse
import hashlib
import json
import sqlite3
import time
from dataclasses import asdict, dataclass
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import TYPE_CHECKING, Any, Final, cast

import httpx
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from stakeholder_intelligence_agent.config import Settings

if TYPE_CHECKING:
    from docx.document import Document as DocumentType

PROJECT_ROOT: Final = Path(__file__).resolve().parents[1]
SEED_REVISION: Final = 3
DOCUMENT_DIRECTORY_NAME: Final = "seed-source-documents"
DOCUMENT_MANIFEST_NAME: Final = "seed-document-manifest.json"
SUMMARY_NAME: Final = "demo-seed-summary.json"
INTERVIEW_TURN_ATTEMPTS: Final = 3
INTERVIEW_TURN_RETRY_SECONDS: Final = 30
INSIGHT_FULL_RUN_RETRY_SECONDS: Final = 65
TERMINAL_INSIGHT_STATUSES: Final = {
    "complete",
    "partial",
    "insufficient_evidence",
    "failed",
}
BLUE: Final = RGBColor(46, 116, 181)
DARK_BLUE: Final = RGBColor(31, 77, 120)
GRAY: Final = RGBColor(89, 89, 89)


@dataclass(frozen=True, slots=True)
class BriefSection:
    """One compact section in a synthetic supporting brief."""

    heading: str
    paragraphs: tuple[str, ...]


@dataclass(frozen=True, slots=True)
class StakeholderSeed:
    """One participant persona, evidence brief, and truthful scripted answers."""

    display_name: str
    role: str
    department: str
    document_filename: str
    document_title: str
    document_subtitle: str
    document_sections: tuple[BriefSection, ...]
    answers: tuple[str, ...]
    final_check_answer: str
    expire_after_answers: int | None = None


@dataclass(frozen=True, slots=True)
class EngagementSeed:
    """One engagement with participants and PM insight questions."""

    name: str
    description: str
    stakeholders: tuple[StakeholderSeed, ...]
    insight_questions: tuple[str, ...]


ENGAGEMENTS: Final = (
    EngagementSeed(
        name="Northstar Foods Order-to-Cash Stabilization",
        description=(
            "Synthetic discovery engagement for a 12-week phased stabilization of order intake, "
            "credit control, allocation, fulfillment, and customer exception handling."
        ),
        stakeholders=(
            StakeholderSeed(
                display_name="Maya Chen",
                role="Director of Customer Operations",
                department="Customer Operations",
                document_filename="northstar-customer-operations-baseline.docx",
                document_title="Customer Operations Baseline",
                document_subtitle="Northstar Foods order-to-cash stabilization evidence brief",
                document_sections=(
                    BriefSection(
                        "Operating baseline",
                        (
                            "The 18-person Customer Operations team validates EDI and portal "
                            "orders before release to the ERP fulfillment queue.",
                            "Eleven percent of weekly orders require manual correction. During "
                            "promotion weeks the rate rises to approximately 18 percent, mainly "
                            "because of customer master-data, unit-of-measure, and cutoff errors.",
                        ),
                    ),
                    BriefSection(
                        "Ownership and handoffs",
                        (
                            "Commercial owns negotiated terms; Customer Operations validates order "
                            "content; Supply Planning owns constrained allocation; Distribution "
                            "owns pick, pack, and carrier release.",
                            "Substitution approval is not consistently owned when a requested item "
                            "is unavailable after the daily cutoff.",
                        ),
                    ),
                    BriefSection(
                        "Pilot position",
                        (
                            "Customer Operations supports a pilot with the top 20 grocery accounts "
                            "if daily exception huddles, rollback criteria, and account-specific "
                            "training are in place before activation.",
                        ),
                    ),
                ),
                answers=(
                    "I lead 18 customer operations coordinators who validate incoming EDI and portal orders, resolve exceptions, and keep grocery customers informed from order receipt through delivery confirmation.",
                    "An order moves from the customer channel into the ERP validation queue, then to credit review, supply allocation, warehouse release, and carrier booking. My team owns validation and communication, but Finance, Supply Planning, Distribution, and Sales each own a later decision.",
                    "The biggest issue is manual correction: about 11 percent of weekly orders need intervention, rising to roughly 18 percent in promotion weeks. The common causes are customer master-data errors, unit-of-measure mismatches, and orders placed after cutoff.",
                    "Commercial owns negotiated terms, Customer Operations validates the order, Supply Planning decides constrained allocation, and Distribution owns physical release. Product substitution is the unclear handoff, especially after cutoff when Sales has already made a customer promise.",
                    "I support a phased pilot with the top 20 grocery accounts. My conditions are daily exception huddles, explicit rollback criteria, account-specific training, and no expansion until manual correction and on-time confirmation metrics are stable for two weeks.",
                    "Sales sometimes requests a late exception to protect a customer commitment while Finance or Supply Planning applies the standard rule. We need one documented exception authority and a response deadline, not negotiation in email on every order.",
                    "The executive sponsor should resolve the substitution owner and approve the pilot guardrails. We should track manual correction rate, time to confirm, on-time-in-full delivery, credit holds, and customer escalations by account.",
                ),
                final_check_answer=(
                    "Nothing else is material. The unresolved substitution owner and the need for "
                    "account-level rollback criteria are the points I would want recorded."
                ),
            ),
            StakeholderSeed(
                display_name="Daniel Ortiz",
                role="Finance Transformation Controller",
                department="Finance",
                document_filename="northstar-finance-control-readiness.docx",
                document_title="Finance Control Readiness Memo",
                document_subtitle="Northstar Foods order-to-cash stabilization evidence brief",
                document_sections=(
                    BriefSection(
                        "Control baseline",
                        (
                            "Finance reviews credit holds, price overrides, tax exceptions, credit "
                            "notes, and daily revenue-interface reconciliation.",
                            "Exception reconciliation currently closes in a median 2.4 business "
                            "days; the target for the pilot is one business day.",
                        ),
                    ),
                    BriefSection(
                        "Material risks",
                        (
                            "Late commercial overrides can bypass documented approval evidence and "
                            "create mismatches between customer terms, invoices, and revenue records.",
                            "A shared exception reason code and immutable approval record are "
                            "required for auditability.",
                        ),
                    ),
                    BriefSection(
                        "Pilot position",
                        (
                            "Finance gives conditional support if segregation of duties, named "
                            "override authority, daily reconciliation, and rollback thresholds are "
                            "tested before the first account is activated.",
                        ),
                    ),
                ),
                answers=(
                    "I own the financial control design across credit holds, price overrides, tax exceptions, credit notes, and the daily reconciliation between order management, invoicing, and the revenue interface.",
                    "Customer Operations submits an exception with the order evidence, Credit reviews exposure, and Finance approves any price or term override. Once goods ship, Billing creates the invoice and my team reconciles the revenue interface and unresolved credit notes.",
                    "The median exception reconciliation takes about 2.4 business days. Late commercial overrides and inconsistent reason codes are the main causes, and both make it difficult to prove who approved the final financial outcome.",
                    "My largest risk is an override being accepted after warehouse release without a durable approval record. That can produce invoice disputes, revenue mismatch, and weak segregation of duties even when the customer outcome looks successful.",
                    "I conditionally support the 12-week pilot if named override authority, immutable approval evidence, daily reconciliation, and rollback thresholds are tested before the first account goes live. I would not support an informal waiver process.",
                    "Customer Operations wants speed and Sales wants flexibility, while Finance needs control evidence. The compromise is a time-boxed exception route with pre-approved categories, not a blanket removal of credit or price controls.",
                    "The steering group still needs to define the one-business-day reconciliation owner and decide whether unresolved exceptions pause expansion. Those decisions should be written into the pilot acceptance criteria.",
                ),
                final_check_answer=(
                    "No additional issue is more important than preserving the approval record and "
                    "making unresolved reconciliation a formal expansion gate."
                ),
            ),
            StakeholderSeed(
                display_name="Priya Nair",
                role="Regional Sales Director",
                department="Commercial",
                document_filename="northstar-commercial-customer-risk-note.docx",
                document_title="Commercial Customer Risk Note",
                document_subtitle="Northstar Foods order-to-cash stabilization evidence brief",
                document_sections=(
                    BriefSection(
                        "Customer commitments",
                        (
                            "The region manages 42 grocery accounts; the proposed pilot group of 20 "
                            "accounts represents 61 percent of regional revenue.",
                            "Six pilot candidates use promotion calendars that change inside the "
                            "standard order cutoff and therefore depend on a bounded exception path.",
                        ),
                    ),
                    BriefSection(
                        "Commercial concerns",
                        (
                            "Removing local override flexibility without a response service level "
                            "could delay confirmation and damage customer trust.",
                            "Account teams need advance messaging, named escalation contacts, and a "
                            "recorded customer-impact review before each activation wave.",
                        ),
                    ),
                    BriefSection(
                        "Pilot position",
                        (
                            "Commercial supports a smaller first wave of eight accounts, followed by "
                            "expansion only after confirmation time and escalation trends are reviewed.",
                        ),
                    ),
                ),
                answers=(
                    "I lead the regional account directors for 42 grocery customers and own commercial commitments, promotion planning, escalation handling, and coordination with Customer Operations when an order cannot be confirmed as promised.",
                    "Account teams capture negotiated terms and promotion calendars, Customer Operations validates the order, and Supply Planning confirms availability. When a promotion changes after cutoff, my team asks for an exception and needs a fast decision from Operations and Finance.",
                    "The proposed 20-account pilot covers about 61 percent of regional revenue. Six of those accounts frequently change promotion quantities inside the normal cutoff, so a rigid no-exception design would create avoidable confirmation delays.",
                    "I am concerned that control improvements could remove local flexibility without adding a response service level. A late answer is often as damaging as a rejection because the customer cannot re-plan its promotion or shelf allocation.",
                    "I support a pilot, but I prefer an eight-account first wave with advance customer messaging, named escalation contacts, and a review of confirmation time and customer complaints before expanding to the remaining accounts.",
                    "Finance is right that approval evidence matters, but Commercial needs pre-approved exception categories and a decision deadline. I disagree with pausing every order that has a minor documentation gap after the customer commitment is already confirmed.",
                    "The program manager should segment accounts by promotion volatility rather than revenue alone and should include customer-impact evidence in the expansion decision, not just internal processing metrics.",
                ),
                final_check_answer=(
                    "Nothing further. The account segmentation and a measurable response deadline "
                    "for exceptions are the two commercial conditions I want retained."
                ),
            ),
        ),
        insight_questions=(
            "What responsibilities, handoffs, and operational risks are supported for the Northstar Foods order-to-cash stabilization?",
            "Where do the Northstar Foods stakeholders agree or disagree on the phased rollout, and what conditions shape their support?",
            "What evidence gaps and follow-up actions should the program manager address before approving the Northstar Foods pilot?",
        ),
    ),
    EngagementSeed(
        name="Alderon Manufacturing Predictive Maintenance Pilot",
        description=(
            "Synthetic discovery engagement for a 90-day predictive-maintenance pilot across three "
            "plants, focused on maintenance workflow, production impact, and OT risk readiness."
        ),
        stakeholders=(
            StakeholderSeed(
                display_name="Liam Brooks",
                role="Plant Operations Manager",
                department="Operations",
                document_filename="alderon-plant-operations-baseline.docx",
                document_title="Plant Operations Baseline",
                document_subtitle="Alderon Manufacturing predictive-maintenance pilot evidence brief",
                document_sections=(
                    BriefSection(
                        "Production baseline",
                        (
                            "The pilot line experienced 37 hours of unplanned downtime in the last "
                            "quarter. Eleven hours were linked to repeat bearing and motor issues.",
                            "Production supervisors currently receive maintenance updates through "
                            "radio calls and end-of-shift notes rather than one shared workflow.",
                        ),
                    ),
                    BriefSection(
                        "Decision needs",
                        (
                            "Operations needs a named owner for alert triage, a response target by "
                            "severity, and a rule for when production may defer an inspection.",
                            "The pilot must not stop a line automatically; a supervisor and a "
                            "maintenance lead must review the evidence together.",
                        ),
                    ),
                    BriefSection(
                        "Pilot position",
                        (
                            "Operations supports one-line deployment for 90 days if alerts are "
                            "advisory, planned downtime is protected, and false alerts are reviewed weekly.",
                        ),
                    ),
                ),
                answers=(
                    "I am accountable for safe production, shift performance, schedule recovery, and coordination with Maintenance when equipment condition threatens output on the pilot line.",
                    "Operators report abnormal noise or heat to the shift supervisor, who calls Maintenance. The maintenance planner then decides whether to inspect immediately or place work in the next planned window, but the decision is often only captured in radio traffic and shift notes.",
                    "The pilot line had 37 hours of unplanned downtime last quarter, including 11 hours from repeat bearing and motor issues. We also lose time when teams debate whether an alert justifies stopping a line.",
                    "Operations should own the production-impact decision and Maintenance should own the technical diagnosis. The missing role is a named alert-triage owner who combines the sensor evidence with line context and records the decision.",
                    "I support a 90-day one-line pilot if alerts remain advisory, planned downtime is protected, and a supervisor and maintenance lead jointly approve any interruption. False alerts must be reviewed every week.",
                    "My concern is that the project could optimize model accuracy while ignoring schedule recovery and operator trust. A technically valid alert still needs a response target and a practical maintenance window.",
                    "Success should include avoided downtime, repeat-failure reduction, response time by severity, false-alert review, and whether operators received a clear disposition before shift handover.",
                ),
                final_check_answer=(
                    "No. The important condition is that the pilot remains advisory and every alert "
                    "gets a recorded operational disposition before the next shift."
                ),
            ),
            StakeholderSeed(
                display_name="Sofia Petrova",
                role="Maintenance Reliability Lead",
                department="Engineering",
                document_filename="alderon-maintenance-reliability-readiness.docx",
                document_title="Maintenance Reliability Readiness",
                document_subtitle="Alderon Manufacturing predictive-maintenance pilot evidence brief",
                document_sections=(
                    BriefSection(
                        "Maintenance baseline",
                        (
                            "The maintenance team receives approximately 24 condition-related work "
                            "requests per week; nine arrive without a consistent severity or asset code.",
                            "Repeat bearing and motor faults are concentrated on seven assets that "
                            "already have usable maintenance history.",
                        ),
                    ),
                    BriefSection(
                        "Workflow proposal",
                        (
                            "Reliability Engineering proposes one alert queue, four severity levels, "
                            "and a named duty engineer who records disposition and linked work order.",
                            "Calibration and false-positive review should occur weekly with Operations.",
                        ),
                    ),
                    BriefSection(
                        "Pilot position",
                        (
                            "Engineering supports the pilot on seven assets if sensor quality checks, "
                            "protected inspection time, and CMMS integration ownership are confirmed.",
                        ),
                    ),
                ),
                answers=(
                    "I own the reliability strategy, condition-monitoring standards, failure analysis, and the maintenance workflow that turns an observed condition into an inspection, work order, repair, and learning record.",
                    "We receive around 24 condition-related requests each week, and about nine arrive without a consistent severity or asset code. The planner has to reconstruct context before a technician can even decide whether the issue is urgent.",
                    "Seven assets account for most repeat bearing and motor events and already have useful maintenance history. Those assets are the strongest pilot scope because we can compare alerts with known failures and completed work orders.",
                    "Reliability Engineering should own alert criteria and technical disposition, the duty engineer should triage each alert, Maintenance Planning should create the work order, and Operations should approve the production window.",
                    "I support the pilot if we use one queue, four severity levels, a named duty engineer, weekly calibration with Operations, and a required link from every accepted alert to its CMMS work order.",
                    "I disagree with treating every alert as merely advisory forever. During the pilot that is sensible, but a validated critical condition eventually needs a predefined escalation rule rather than optional review.",
                    "We still need ownership for sensor health and CMMS integration failures. Without those controls, the team could mistake missing or stale data for healthy equipment and overstate pilot performance.",
                ),
                final_check_answer=(
                    "Nothing else. Sensor-health ownership and traceability from alert to work order "
                    "are the evidence conditions I would not compromise."
                ),
            ),
            StakeholderSeed(
                display_name="Noah Williams",
                role="OT Cybersecurity Manager",
                department="Technology Risk",
                document_filename="alderon-ot-security-review-draft.docx",
                document_title="OT Security Review Draft",
                document_subtitle="Alderon Manufacturing predictive-maintenance pilot evidence brief",
                document_sections=(
                    BriefSection(
                        "Proposed data path",
                        (
                            "The draft architecture sends sensor readings from the line gateway to "
                            "an analytics service through the plant OT demilitarized zone.",
                            "Vendor remote support is proposed but its access window, approval owner, "
                            "session recording, and emergency revocation path are not yet documented.",
                        ),
                    ),
                    BriefSection(
                        "Open controls",
                        (
                            "Asset inventory validation, gateway hardening, certificate rotation, "
                            "log retention, and incident ownership require review before connection.",
                            "The review is explicitly a draft and does not represent security approval.",
                        ),
                    ),
                    BriefSection(
                        "Review status",
                        (
                            "Technology Risk has not completed the interview or approved the pilot "
                            "architecture. A replacement invitation is required after session expiry.",
                        ),
                    ),
                ),
                answers=(
                    "I review the OT network boundary, gateway hardening, identity and certificate controls, vendor remote access, logging, and the incident response ownership for any new plant-connected service.",
                    "The current draft sends sensor readings through a line gateway and the OT demilitarized zone to an analytics service. Vendor support is mentioned, but the access window, approval owner, session recording, and emergency revocation path are not documented.",
                    "Before connection I need validated asset inventory, an approved data-flow diagram, gateway configuration evidence, certificate rotation ownership, retained security logs, and a tested incident contact path for each plant.",
                ),
                final_check_answer=(
                    "The review is still incomplete, so I cannot provide a final security position."
                ),
                expire_after_answers=2,
            ),
        ),
        insight_questions=(
            "What responsibilities, handoffs, and operational risks are supported for the Alderon Manufacturing predictive-maintenance pilot?",
            "Where do completed Alderon interviews and documents show support, concern, or disagreement about the 90-day pilot?",
            "What readiness evidence is missing before Alderon connects predictive-maintenance technology to the plant network?",
        ),
    ),
)


class SeedError(RuntimeError):
    """Represent a safe, credential-free seed failure."""


def _set_font(run: Any, *, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def _configure_document(document: DocumentType) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(16)
    heading.font.color.rgb = BLUE
    heading.font.bold = True
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    _set_font(
        header.add_run("Synthetic stakeholder evidence | Demonstration only"),
        size=8.5,
        color=GRAY,
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    _set_font(
        footer.add_run("Prepared 20 July 2026 | Not a production business record"),
        size=8.5,
        color=GRAY,
    )


def build_supporting_document(seed: StakeholderSeed, destination: Path) -> None:
    """Create one compact, explicitly synthetic stakeholder evidence brief."""
    document = Document()
    _configure_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_font(title.add_run(seed.document_title), size=24, color=DARK_BLUE, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_font(subtitle.add_run(seed.document_subtitle), size=12, color=GRAY)

    for label, value in (
        ("Prepared by", seed.display_name),
        ("Role", seed.role),
        ("Department", seed.department),
        ("Evidence status", "Synthetic interview supporting material"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _set_font(paragraph.add_run(f"{label}: "), size=10.5, color=GRAY, bold=True)
        _set_font(paragraph.add_run(value), size=10.5, color=GRAY)

    for section in seed.document_sections:
        document.add_heading(section.heading, level=1)
        for text in section.paragraphs:
            paragraph = document.add_paragraph(text)
            paragraph.paragraph_format.widow_control = True

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    _set_font(note.add_run("Source note: "), size=9, color=GRAY, bold=True)
    _set_font(
        note.add_run(
            "This fictional brief was created solely to exercise the local Stakeholder "
            "Intelligence Agent with realistic, non-confidential evidence."
        ),
        size=9,
        color=GRAY,
    )

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))


def _file_sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _document_seed_fingerprint(seed: StakeholderSeed) -> str:
    payload = {
        "seed_revision": SEED_REVISION,
        "stakeholder": asdict(seed),
    }
    encoded = json.dumps(payload, sort_keys=True, separators=(",", ":")).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def _load_json_object(path: Path) -> dict[str, Any]:
    if not path.is_file():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return payload if isinstance(payload, dict) else {}


def prepare_documents(
    document_root: Path,
    manifest_path: Path,
) -> tuple[Path, ...]:
    """Create or reuse revision-matched briefs and persist their exact hashes."""
    previous = _load_json_object(manifest_path)
    previous_documents = previous.get("documents", {})
    if previous.get("seed_revision") != SEED_REVISION or not isinstance(previous_documents, dict):
        previous_documents = {}

    paths: list[Path] = []
    manifest_documents: dict[str, dict[str, str]] = {}
    for engagement in ENGAGEMENTS:
        for stakeholder in engagement.stakeholders:
            destination = document_root / stakeholder.document_filename
            fingerprint = _document_seed_fingerprint(stakeholder)
            record = previous_documents.get(stakeholder.document_filename, {})
            reusable = (
                isinstance(record, dict)
                and record.get("source_fingerprint") == fingerprint
                and destination.is_file()
                and record.get("content_sha256") == _file_sha256(destination)
            )
            if not reusable:
                build_supporting_document(stakeholder, destination)
            content_hash = _file_sha256(destination)
            manifest_documents[stakeholder.document_filename] = {
                "source_fingerprint": fingerprint,
                "content_sha256": content_hash,
            }
            paths.append(destination)

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        json.dumps(
            {
                "seed_revision": SEED_REVISION,
                "documents": manifest_documents,
            },
            indent=2,
            sort_keys=True,
        )
        + "\n",
        encoding="utf-8",
    )
    return tuple(paths)


class LiveApi:
    """Small credential-safe client for the approved domain routes."""

    def __init__(self, base_url: str, settings: Settings) -> None:
        self._settings = settings
        self._document_root = settings.data_root / DOCUMENT_DIRECTORY_NAME
        self._client = httpx.Client(
            base_url=base_url,
            timeout=httpx.Timeout(900.0, connect=10.0),
            headers={"X-Correlation-ID": "demo-seed"},
        )
        self._pm_token = self._activate_pm()

    def close(self) -> None:
        self._client.close()

    def _activate_pm(self) -> str:
        response = self._client.post(
            "/api/v1/auth/pm/activate",
            json={
                "bootstrap_token": self._settings.pm_bootstrap_token.get_secret_value(),
            },
        )
        payload = self._json(response, expected={200}, operation="PM activation")
        return str(payload["access_token"])

    def _pm_headers(self) -> dict[str, str]:
        return {"Authorization": f"Bearer {self._pm_token}"}

    @staticmethod
    def _stakeholder_headers(token: str) -> dict[str, str]:
        return {"Authorization": f"Bearer {token}"}

    @staticmethod
    def _json(
        response: httpx.Response,
        *,
        expected: set[int],
        operation: str,
    ) -> dict[str, Any]:
        if response.status_code not in expected:
            try:
                error = response.json().get("error", {})
                code = error.get("code", "HTTP_ERROR")
                message = error.get("message", "The operation failed.")
            except (ValueError, AttributeError):
                code = "HTTP_ERROR"
                message = "The operation returned a non-JSON failure."
            raise SeedError(f"{operation} failed ({response.status_code}, {code}): {message}")
        payload = response.json()
        if not isinstance(payload, dict):
            raise SeedError(f"{operation} returned an invalid response envelope.")
        return payload

    def engagement(self, seed: EngagementSeed) -> dict[str, Any]:
        listed = self._json(
            self._client.get("/api/v1/pm/engagements", headers=self._pm_headers()),
            expected={200},
            operation="Engagement list",
        )["engagements"]
        matches = [item for item in listed if item["name"] == seed.name]
        if len(matches) > 1:
            raise SeedError(f"Multiple engagements have the reserved demo name {seed.name!r}.")
        if matches:
            engagement = matches[0]
            self._json(
                self._client.post(
                    f"/api/v1/pm/engagements/{engagement['engagement_id']}/select",
                    headers=self._pm_headers(),
                ),
                expected={200},
                operation="Engagement selection",
            )
            return cast("dict[str, Any]", engagement)
        payload = self._json(
            self._client.post(
                "/api/v1/pm/engagements",
                json={"name": seed.name, "description": seed.description},
                headers=self._pm_headers(),
            ),
            expected={201},
            operation="Engagement creation",
        )
        return cast("dict[str, Any]", payload["engagement"])

    def stakeholder(
        self,
        engagement_id: str,
        seed: StakeholderSeed,
    ) -> dict[str, Any]:
        path = f"/api/v1/pm/engagements/{engagement_id}/stakeholders"
        listed = self._json(
            self._client.get(path, headers=self._pm_headers()),
            expected={200},
            operation="Stakeholder list",
        )["stakeholders"]
        matches = [item for item in listed if item["display_name"] == seed.display_name]
        if len(matches) > 1:
            raise SeedError(f"Duplicate demo stakeholder {seed.display_name!r}.")
        if matches:
            return cast("dict[str, Any]", matches[0])
        created = self._json(
            self._client.post(
                path,
                json={
                    "display_name": seed.display_name,
                    "role": seed.role,
                    "department": seed.department,
                },
                headers=self._pm_headers(),
            ),
            expected={201},
            operation="Stakeholder creation",
        )["stakeholder"]
        return cast("dict[str, Any]", created)

    def _interviews(self, engagement_id: str) -> list[dict[str, Any]]:
        payload = self._json(
            self._client.get(
                f"/api/v1/pm/engagements/{engagement_id}/interviews",
                headers=self._pm_headers(),
            ),
            expected={200},
            operation="Interview list",
        )
        return list(payload["interview_sessions"])

    def _invitations(self, engagement_id: str) -> list[dict[str, Any]]:
        payload = self._json(
            self._client.get(
                f"/api/v1/pm/engagements/{engagement_id}/invitations",
                headers=self._pm_headers(),
            ),
            expected={200},
            operation="Invitation list",
        )
        return list(payload["invitations"])

    def interview_for(self, engagement_id: str, stakeholder_id: str) -> dict[str, Any] | None:
        matches = [
            item
            for item in self._interviews(engagement_id)
            if item["stakeholder_id"] == stakeholder_id
        ]
        if len(matches) > 1:
            raise SeedError("A demo stakeholder unexpectedly has multiple interview sessions.")
        return matches[0] if matches else None

    def invitation_token(self, engagement_id: str, stakeholder_id: str) -> str:
        candidates = [
            item
            for item in self._invitations(engagement_id)
            if item["stakeholder_id"] == stakeholder_id
            and item["status"] in {"active", "activated"}
        ]
        if candidates:
            current = candidates[-1]
            response = self._client.get(
                (
                    f"/api/v1/pm/engagements/{engagement_id}/invitations/"
                    f"{current['invitation_id']}/link"
                ),
                headers=self._pm_headers(),
            )
            if response.status_code == 200:
                return str(response.json()["invitation_token"])
            if response.status_code != 409:
                self._json(response, expected={200}, operation="Invitation recovery")

        payload = self._json(
            self._client.post(
                (
                    f"/api/v1/pm/engagements/{engagement_id}/stakeholders/"
                    f"{stakeholder_id}/invitations"
                ),
                headers=self._pm_headers(),
            ),
            expected={201},
            operation="Invitation issuance",
        )
        return str(payload["invitation_token"])

    def activate_stakeholder(self, invitation_token: str) -> tuple[str, dict[str, Any]]:
        payload = self._json(
            self._client.post(
                "/api/v1/auth/stakeholder/activate",
                json={"invitation_token": invitation_token},
            ),
            expected={200},
            operation="Stakeholder activation",
        )
        return str(payload["session"]["access_token"]), payload

    def ensure_document(self, stakeholder_token: str, seed: StakeholderSeed) -> dict[str, Any]:
        headers = self._stakeholder_headers(stakeholder_token)
        listed = self._json(
            self._client.get("/api/v1/stakeholder/documents", headers=headers),
            expected={200},
            operation="Stakeholder document list",
        )["documents"]
        matches = [
            item for item in listed if item["source"]["original_filename"] == seed.document_filename
        ]
        document_path = self._document_root / seed.document_filename
        local_hash = _file_sha256(document_path)
        matching_versions = [
            item
            for item in matches
            if item["latest_version"]["content_hash"] == local_hash
            and item["latest_version"]["state"] == "READY"
        ]
        if matching_versions:
            return cast("dict[str, Any]", matching_versions[-1])
        with document_path.open("rb") as stream:
            response = self._client.post(
                "/api/v1/stakeholder/documents",
                files={
                    "upload": (
                        seed.document_filename,
                        stream,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
                headers=headers,
            )
        uploaded = self._json(
            response,
            expected={201},
            operation="Stakeholder document ingestion",
        )["document"]
        if uploaded["latest_version"]["content_hash"] != local_hash:
            raise SeedError("The ingested document hash does not match the generated source bytes.")
        if uploaded["latest_version"]["state"] != "READY":
            raise SeedError("The ingested seed document did not reach READY.")
        return cast("dict[str, Any]", uploaded)

    def verify_existing_document(
        self,
        engagement_id: str,
        stakeholder_id: str,
        seed: StakeholderSeed,
    ) -> dict[str, Any]:
        """Verify an immutable completed/expired interview's document by filename and hash."""
        documents = self._json(
            self._client.get(
                f"/api/v1/pm/engagements/{engagement_id}/documents",
                headers=self._pm_headers(),
            ),
            expected={200},
            operation="PM document list",
        )["documents"]
        local_hash = _file_sha256(self._document_root / seed.document_filename)
        matches = [
            item
            for item in documents
            if item["source"]["stakeholder_id"] == stakeholder_id
            and item["source"]["original_filename"] == seed.document_filename
            and item["latest_version"]["content_hash"] == local_hash
            and item["latest_version"]["state"] == "READY"
        ]
        if len(matches) != 1:
            raise SeedError(
                f"Expected one READY hash-matched document for {seed.display_name!r}; "
                "the existing seed cannot be reused safely."
            )
        return cast("dict[str, Any]", matches[0])

    def verify_retained_draft_document(
        self,
        engagement_id: str,
        stakeholder_id: str,
        seed: StakeholderSeed,
    ) -> dict[str, Any]:
        """Verify READY lineage hidden from PM inventory while its interview remains draft."""
        local_hash = _file_sha256(self._document_root / seed.document_filename)
        with sqlite3.connect(self._settings.domain_database, timeout=30.0) as connection:
            rows = connection.execute(
                """
                SELECT v.content_hash, v.state
                FROM document_sources AS s
                JOIN document_versions AS v ON v.document_id = s.document_id
                WHERE s.engagement_id = ? AND s.stakeholder_id = ?
                    AND s.original_filename = ? AND s.deleted_at IS NULL
                    AND v.is_active = 1
                ORDER BY v.version_number DESC
                """,
                (engagement_id, stakeholder_id, seed.document_filename),
            ).fetchall()
        if rows != [(local_hash, "READY")]:
            raise SeedError(
                f"Expected one READY hash-matched retained draft document for "
                f"{seed.display_name!r}; the existing seed cannot be reused safely."
            )
        return {"latest_version": {"content_hash": local_hash, "state": "READY"}}

    def interview_status(self, stakeholder_token: str) -> dict[str, Any]:
        return self._json(
            self._client.get(
                "/api/v1/stakeholder/interview/status",
                headers=self._stakeholder_headers(stakeholder_token),
            ),
            expected={200},
            operation="Interview status",
        )

    def start_interview(self, stakeholder_token: str) -> dict[str, Any]:
        return self._json(
            self._client.post(
                "/api/v1/stakeholder/interview/start",
                headers=self._stakeholder_headers(stakeholder_token),
            ),
            expected={200},
            operation="Interview start",
        )

    def submit_answer(self, stakeholder_token: str, answer: str, message_id: str) -> None:
        for attempt in range(1, INTERVIEW_TURN_ATTEMPTS + 1):
            response = self._client.post(
                "/api/v1/stakeholder/interview/turns/stream",
                json={"original_text": answer, "message_id": message_id},
                headers=self._stakeholder_headers(stakeholder_token),
            )
            if response.status_code != 200:
                self._json(response, expected={200}, operation="Interview turn")
            failed = "event: failure" in response.text
            completed = "event: message" in response.text
            if not failed and completed:
                return
            if attempt == INTERVIEW_TURN_ATTEMPTS:
                if failed:
                    raise SeedError("The interview agent emitted a safe failure event.")
                raise SeedError("The interview agent did not emit a completed assistant turn.")
            delay = INTERVIEW_TURN_RETRY_SECONDS * attempt
            print(
                f"      transient interview turn failure; retrying same message in {delay}s",
                flush=True,
            )
            time.sleep(delay)
        raise SeedError("The interview turn retry loop ended without a terminal result.")

    def finish_interview(self, stakeholder_token: str) -> dict[str, Any]:
        return self._json(
            self._client.post(
                "/api/v1/stakeholder/interview/finish",
                headers=self._stakeholder_headers(stakeholder_token),
            ),
            expected={200},
            operation="Interview finalization",
        )

    def expire_interview(
        self,
        *,
        interview_session_id: str,
    ) -> None:
        """Simulate elapsed wall time, then let AccessService persist expiration."""
        expired_at = (datetime.now(UTC) - timedelta(seconds=1)).isoformat().replace("+00:00", "Z")
        with sqlite3.connect(self._settings.domain_database, timeout=30.0) as connection:
            row = connection.execute(
                """
                SELECT invitation_id, engagement_id FROM interview_sessions
                WHERE interview_session_id = ?
                """,
                (interview_session_id,),
            ).fetchone()
            if row is None:
                raise SeedError("The interview selected for expiration does not exist.")
            invitation_id = str(row[0])
            engagement_id = str(row[1])
            connection.execute(
                "UPDATE invitation_tokens SET expires_at = ? WHERE invitation_id = ?",
                (expired_at, invitation_id),
            )
            connection.execute(
                """
                UPDATE access_sessions SET expires_at = ?
                WHERE interview_session_id = ? AND principal_type = 'stakeholder'
                """,
                (expired_at, interview_session_id),
            )
        response = self._client.get(
            (f"/api/v1/pm/engagements/{engagement_id}/invitations/{invitation_id}/link"),
            headers=self._pm_headers(),
        )
        if response.status_code != 409:
            raise SeedError("AccessService did not reject recovery of the expired invitation.")
        with sqlite3.connect(self._settings.domain_database, timeout=30.0) as connection:
            status = connection.execute(
                "SELECT status FROM invitation_tokens WHERE invitation_id = ?",
                (invitation_id,),
            ).fetchone()
        if status is None or status[0] != "expired":
            raise SeedError("AccessService did not persist the expected expired status.")

    def _reusable_insight(
        self,
        engagement_id: str,
        question: str,
        reusable_run_id: str | None,
        existing: list[dict[str, Any]],
    ) -> dict[str, Any] | None:
        """Find a validated completed run even when a prior seed stopped before its summary."""
        reusable = [
            item
            for item in existing
            if item["requested_question"] == question
            and item["status"] in {"complete", "partial", "insufficient_evidence"}
        ]
        if reusable_run_id is not None:
            reusable.sort(key=lambda item: item["run_id"] != reusable_run_id)
        for candidate in reusable:
            candidate_run_id = str(candidate["run_id"])
            report_response = self._client.get(
                (f"/api/v1/pm/engagements/{engagement_id}/insights/{candidate_run_id}/report"),
                headers=self._pm_headers(),
            )
            if report_response.status_code == 200:
                envelope = report_response.json()
                report = envelope.get("report", {}) if isinstance(envelope, dict) else {}
                if isinstance(report, dict) and report.get("citations"):
                    return cast(
                        "dict[str, Any]",
                        candidate
                        | {
                            "citation_count": len(report["citations"]),
                            "metrics": envelope.get("metrics", {}),
                        },
                    )
            print(
                f"      stored run {candidate_run_id} failed current report validation; rerunning",
                flush=True,
            )
        return None

    def run_insight(
        self,
        engagement_id: str,
        question: str,
        reusable_run_id: str | None,
    ) -> dict[str, Any]:
        existing = cast(
            "list[dict[str, Any]]",
            self._json(
                self._client.get(
                    f"/api/v1/pm/engagements/{engagement_id}/insights",
                    headers=self._pm_headers(),
                ),
                expected={200},
                operation="Insight history",
            )["runs"],
        )
        reusable = self._reusable_insight(
            engagement_id,
            question,
            reusable_run_id,
            existing,
        )
        if reusable is not None:
            return reusable

        failure_codes: list[str] = []
        for attempt in range(1, 4):
            created = self._json(
                self._client.post(
                    f"/api/v1/pm/engagements/{engagement_id}/insights",
                    json={"question": question},
                    headers=self._pm_headers(),
                ),
                expected={202},
                operation="Insight start",
            )["run"]
            run_id = str(created["run_id"])
            last_status = ""
            deadline = time.monotonic() + 1_200
            while time.monotonic() < deadline:
                try:
                    response = self._client.get(
                        f"/api/v1/pm/engagements/{engagement_id}/insights/{run_id}",
                        headers=self._pm_headers(),
                    )
                except httpx.TransportError:
                    print(
                        "      transient Agent Server connection loss; reconnecting",
                        flush=True,
                    )
                    time.sleep(5)
                    continue
                current = self._json(
                    response,
                    expected={200},
                    operation="Insight status",
                )["run"]
                status = str(current["status"])
                if status != last_status:
                    print(f"      insight {run_id}: {status}", flush=True)
                    last_status = status
                if status not in TERMINAL_INSIGHT_STATUSES:
                    time.sleep(5)
                    continue
                if status == "failed":
                    code = str(current.get("failure_code") or "INSIGHT_EXECUTION_FAILED")
                    failure_codes.append(code)
                    if attempt < 3:
                        print(
                            "      retrying the full research run after safe failure "
                            f"{code} in {INSIGHT_FULL_RUN_RETRY_SECONDS}s",
                            flush=True,
                        )
                        time.sleep(INSIGHT_FULL_RUN_RETRY_SECONDS)
                        break
                    raise SeedError(
                        "Live insight exhausted three full-run attempts: "
                        + ", ".join(failure_codes)
                    )
                report_envelope = self._json(
                    self._client.get(
                        (f"/api/v1/pm/engagements/{engagement_id}/insights/{run_id}/report"),
                        headers=self._pm_headers(),
                    ),
                    expected={200},
                    operation="Insight report",
                )
                report = report_envelope["report"]
                if not report["citations"]:
                    raise SeedError("The live insight report has no validated citations.")
                return cast(
                    "dict[str, Any]",
                    current
                    | {
                        "citation_count": len(report["citations"]),
                        "metrics": report_envelope["metrics"],
                    },
                )
            else:
                raise SeedError("The live insight run exceeded the 20-minute seed timeout.")
        raise SeedError("The live insight retry loop ended without a terminal result.")


def _stakeholder_turn_count(status: dict[str, Any]) -> int:
    return sum(1 for turn in status["turns"] if turn["speaker"] == "stakeholder")


def _last_assistant_text(status: dict[str, Any]) -> str:
    for turn in reversed(status["turns"]):
        if turn["speaker"] == "assistant":
            return str(turn["text"])
    return ""


def _next_answer(seed: StakeholderSeed, status: dict[str, Any]) -> str:
    if "Before we finish, is there anything important" in _last_assistant_text(status):
        return seed.final_check_answer
    index = _stakeholder_turn_count(status)
    if index < len(seed.answers):
        return seed.answers[index]
    return seed.final_check_answer


def seed_interview(
    api: LiveApi,
    engagement_id: str,
    stakeholder: dict[str, Any],
    seed: StakeholderSeed,
) -> dict[str, Any]:
    """Run or resume one real interview and reach its intended lifecycle state."""
    stakeholder_id = str(stakeholder["stakeholder_id"])
    current = api.interview_for(engagement_id, stakeholder_id)
    if current is not None and current["status"] == "ready":
        document = api.verify_existing_document(
            engagement_id,
            stakeholder_id,
            seed,
        )
        print(f"    {seed.display_name}: already completed", flush=True)
        return {
            "status": "ready",
            "interview_session_id": current["interview_session_id"],
            "document_content_hash": document["latest_version"]["content_hash"],
        }
    invitations = [
        item for item in api._invitations(engagement_id) if item["stakeholder_id"] == stakeholder_id
    ]
    if (
        seed.expire_after_answers is not None
        and current is not None
        and invitations
        and invitations[-1]["status"] == "expired"
    ):
        document = api.verify_retained_draft_document(
            engagement_id,
            stakeholder_id,
            seed,
        )
        print(f"    {seed.display_name}: already expired with a retained draft", flush=True)
        return {
            "status": "expired",
            "interview_session_id": current["interview_session_id"],
            "document_content_hash": document["latest_version"]["content_hash"],
        }

    invitation_token = api.invitation_token(engagement_id, stakeholder_id)
    stakeholder_token, activation = api.activate_stakeholder(invitation_token)
    interview_session_id = str(activation["interview_session_id"])
    document = api.ensure_document(stakeholder_token, seed)
    print(
        f"    {seed.display_name}: document {document['latest_version']['state']}",
        flush=True,
    )
    status = api.start_interview(stakeholder_token)
    print(f"    {seed.display_name}: live interview started", flush=True)

    while True:
        stakeholder_answers = _stakeholder_turn_count(status)
        if (
            seed.expire_after_answers is not None
            and stakeholder_answers >= seed.expire_after_answers
        ):
            api.expire_interview(
                interview_session_id=interview_session_id,
            )
            print(
                f"    {seed.display_name}: session expired after {stakeholder_answers} answers",
                flush=True,
            )
            return {
                "status": "expired",
                "interview_session_id": interview_session_id,
                "document_content_hash": document["latest_version"]["content_hash"],
            }
        if status["completion_recommended"]:
            finished = api.finish_interview(stakeholder_token)
            if finished["interview_session"]["status"] != "ready":
                raise SeedError("A finalized interview did not reach READY.")
            print(
                f"    {seed.display_name}: interview completed with {stakeholder_answers} answers",
                flush=True,
            )
            return {
                "status": "ready",
                "interview_session_id": interview_session_id,
                "document_content_hash": document["latest_version"]["content_hash"],
            }
        if stakeholder_answers >= 12:
            raise SeedError(
                "The interview assistant did not recommend completion within 12 answers."
            )
        answer = _next_answer(seed, status)
        api.submit_answer(
            stakeholder_token,
            answer,
            message_id=f"demo-{stakeholder_id}-{stakeholder_answers + 1}",
        )
        status = api.interview_status(stakeholder_token)
        print(
            f"      answer {stakeholder_answers + 1}: assistant follow-up saved",
            flush=True,
        )


def _reusable_run_ids(summary: dict[str, Any]) -> dict[tuple[str, str], str]:
    if summary.get("seed_revision") != SEED_REVISION:
        return {}
    reusable: dict[tuple[str, str], str] = {}
    engagements = summary.get("engagements", [])
    if not isinstance(engagements, list):
        return reusable
    for engagement in engagements:
        if not isinstance(engagement, dict) or not isinstance(engagement.get("name"), str):
            continue
        insights = engagement.get("insights", [])
        if not isinstance(insights, list):
            continue
        for insight in insights:
            if (
                isinstance(insight, dict)
                and isinstance(insight.get("question"), str)
                and isinstance(insight.get("run_id"), str)
            ):
                reusable[(engagement["name"], insight["question"])] = insight["run_id"]
    return reusable


def seed_all(base_url: str, settings: Settings) -> dict[str, Any]:
    """Execute the complete non-destructive live seed and return safe identifiers."""
    summary_path = settings.data_root / SUMMARY_NAME
    reusable_runs = _reusable_run_ids(_load_json_object(summary_path))
    api = LiveApi(base_url, settings)
    summary: dict[str, Any] = {
        "generated_at": datetime.now(UTC).isoformat(),
        "seed_revision": SEED_REVISION,
        "synthetic_data": True,
        "research_basis": {
            "classification": "minimum_viable_cross_functional_demo",
            "stakeholders_per_engagement": 3,
            "not_a_saturation_claim": True,
            "benchmark": "GDS guidance normally uses 4-8 participants per research round.",
            "benchmark_url": (
                "https://www.gov.uk/service-manual/user-research/"
                "plan-user-research-for-your-service"
            ),
        },
        "engagements": [],
    }
    try:
        for engagement_seed in ENGAGEMENTS:
            engagement = api.engagement(engagement_seed)
            engagement_id = str(engagement["engagement_id"])
            print(f"  {engagement_seed.name}", flush=True)
            engagement_summary: dict[str, Any] = {
                "engagement_id": engagement_id,
                "name": engagement_seed.name,
                "stakeholders": [],
                "insights": [],
            }
            for stakeholder_seed in engagement_seed.stakeholders:
                stakeholder = api.stakeholder(
                    engagement_id,
                    stakeholder_seed,
                )
                interview = seed_interview(
                    api,
                    engagement_id,
                    stakeholder,
                    stakeholder_seed,
                )
                engagement_summary["stakeholders"].append(
                    {
                        "stakeholder_id": stakeholder["stakeholder_id"],
                        "display_name": stakeholder_seed.display_name,
                        "role": stakeholder_seed.role,
                        "department": stakeholder_seed.department,
                        "document_filename": stakeholder_seed.document_filename,
                        "document_content_hash": interview["document_content_hash"],
                        "interview_session_id": interview["interview_session_id"],
                        "interview_outcome": interview["status"],
                    }
                )
            for index, question in enumerate(engagement_seed.insight_questions, start=1):
                print(
                    f"    research question {index}/{len(engagement_seed.insight_questions)}",
                    flush=True,
                )
                run = api.run_insight(
                    engagement_id,
                    question,
                    reusable_runs.get((engagement_seed.name, question)),
                )
                engagement_summary["insights"].append(
                    {
                        "run_id": run["run_id"],
                        "question": question,
                        "status": run["status"],
                        "citation_count": run["citation_count"],
                        "metrics": run["metrics"],
                    }
                )
            summary["engagements"].append(engagement_summary)
    finally:
        api.close()
    summary_path.write_text(json.dumps(summary, indent=2) + "\n", encoding="utf-8")
    return summary


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--base-url",
        default="http://127.0.0.1:2024",
        help="Loopback Agent Server origin.",
    )
    parser.add_argument(
        "--prepare-documents-only",
        action="store_true",
        help="Generate the six DOCX briefs without mutating application data.",
    )
    args = parser.parse_args()

    settings = Settings()
    document_root = settings.data_root / DOCUMENT_DIRECTORY_NAME
    manifest_path = settings.data_root / DOCUMENT_MANIFEST_NAME
    summary_path = settings.data_root / SUMMARY_NAME
    paths = prepare_documents(document_root, manifest_path)
    print(f"Prepared {len(paths)} synthetic supporting documents.", flush=True)
    if args.prepare_documents_only:
        return 0
    summary = seed_all(str(args.base_url).rstrip("/"), settings)
    print(
        f"Seeded {len(summary['engagements'])} E2E engagements; safe summary: {summary_path}",
        flush=True,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
